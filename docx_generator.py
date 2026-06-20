from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io
import os

TOP_COUNTER_FILE = "top_counter.txt"


def get_next_top_number() -> str:
    if os.path.exists(TOP_COUNTER_FILE):
        with open(TOP_COUNTER_FILE, "r") as f:
            num = int(f.read().strip()) + 1
    else:
        num = 1
    with open(TOP_COUNTER_FILE, "w") as f:
        f.write(str(num))
    return str(num).zfill(2)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:color'), '000000')
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def generate_top_docx(data: dict) -> bytes:
    """Generate the Tender One Pager in MeraPath's exact format"""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.8)
        section.right_margin  = Inches(0.8)

    top_number = get_next_top_number()
    today = datetime.today().strftime("%d %B %Y")

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Tender One Pager (TOP)")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Main table
    table = doc.add_table(rows=12, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def cell_text(cell, text, bold=False, size=10):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def merge_and_set(row_idx, start_col, end_col, text, bold=False, label_only=False):
        cell = table.cell(row_idx, start_col)
        if start_col != end_col:
            cell = cell.merge(table.cell(row_idx, end_col))
        cell_text(cell, text, bold=bold)
        return cell

    # Row 0 — Ref + Date
    r0c0 = table.cell(0, 0).merge(table.cell(0, 2))
    cell_text(r0c0, f"Ref.: MEL/TOP/{top_number}")
    r0c3 = table.cell(0, 3).merge(table.cell(0, 5))
    cell_text(r0c3, f"Date: {today}")

    # Row 1 — Bid Number + RA Option
    cell_text(table.cell(1, 0), "Bid Number:", bold=True)
    r1c1 = table.cell(1, 1).merge(table.cell(1, 2))
    cell_text(r1c1, data.get("bid_number", ""))
    cell_text(table.cell(1, 3), "RA Option (Yes/No):", bold=True)
    r1c4 = table.cell(1, 4).merge(table.cell(1, 5))
    cell_text(r1c4, data.get("ra_option", ""))

    # Row 2 — Start Date + End Date
    cell_text(table.cell(2, 0), "Start Date", bold=True)
    r2c1 = table.cell(2, 1).merge(table.cell(2, 2))
    cell_text(r2c1, data.get("start_date", ""))
    cell_text(table.cell(2, 3), "End Date & Time", bold=True)
    r2c4 = table.cell(2, 4).merge(table.cell(2, 5))
    cell_text(r2c4, data.get("end_date", ""))

    # Row 3 — Department
    r3c0 = table.cell(3, 0).merge(table.cell(3, 5))
    cell_text(r3c0, "Department Name & Address:", bold=True)

    # Row 4 — Department value
    r4c0 = table.cell(4, 0).merge(table.cell(4, 5))
    cell_text(r4c0, data.get("department_name", ""))

    # Row 5 — Estimated Cost + ePBG
    cell_text(table.cell(5, 0), "Estimated Cost", bold=True)
    r5c1 = table.cell(5, 1).merge(table.cell(5, 2))
    cell_text(r5c1, data.get("estimated_cost", ""))
    cell_text(table.cell(5, 3), "ePBG", bold=True)
    r5c4 = table.cell(5, 4).merge(table.cell(5, 5))
    cell_text(r5c4, data.get("epbg", ""))

    # Row 6 — Fees
    r6c0 = table.cell(6, 0).merge(table.cell(6, 1))
    cell_text(r6c0, "Processing Fee (non-refundable):", bold=True)
    cell_text(table.cell(6, 2), data.get("processing_fee", ""))
    r6c3 = table.cell(6, 3).merge(table.cell(6, 4))
    cell_text(r6c3, "Tender Fee (non-refundable)", bold=True)
    cell_text(table.cell(6, 5), data.get("tender_fee", ""))

    # Row 7 — Appraisal fee
    r7c0 = table.cell(7, 0).merge(table.cell(7, 5))
    cell_text(r7c0, f"Appraisal Fee: {data.get('appraisal_fee', 'Not mentioned')}")

    # Row 8 — Project Name label
    r8c0 = table.cell(8, 0).merge(table.cell(8, 5))
    cell_text(r8c0, "Name of Project:", bold=True)

    # Row 9 — Project name value
    r9c0 = table.cell(9, 0).merge(table.cell(9, 5))
    cell_text(r9c0, data.get("project_name", ""))

    # Row 10 — Scope of Work
    r10c0 = table.cell(10, 0).merge(table.cell(10, 5))
    scope_items = data.get("scope_of_work", [])
    r10c0.paragraphs[0].clear()
    p = r10c0.paragraphs[0]
    label_run = p.add_run("Scope of Work:\n")
    label_run.bold = True
    label_run.font.size = Pt(10)
    for i, s in enumerate(scope_items):
        body_run = p.add_run(f"{i+1}. {s}\n")
        body_run.bold = False
        body_run.font.size = Pt(10)

    # Row 11 — Eligibility
    r11c0 = table.cell(11, 0).merge(table.cell(11, 5))
    elig_items = data.get("eligibility_criteria", [])
    r11c0.paragraphs[0].clear()
    p = r11c0.paragraphs[0]
    label_run = p.add_run("Minimum Eligibility Criteria:\n")
    label_run.bold = True
    label_run.font.size = Pt(10)
    for i, e in enumerate(elig_items):
        body_run = p.add_run(f"{i+1}. {e}\n")
        body_run.bold = False
        body_run.font.size = Pt(10)

    # Add extra rows for docs + remarks
    doc_row = table.add_row()
    doc_cell = doc_row.cells[0].merge(doc_row.cells[5])
    doc_items = data.get("documents_required", [])
    doc_cell.paragraphs[0].clear()
    p = doc_cell.paragraphs[0]
    label_run = p.add_run("Documents what we need to focus:\n")
    label_run.bold = True
    label_run.font.size = Pt(10)
    for i, d in enumerate(doc_items):
        body_run = p.add_run(f"{i+1}. {d}\n")
        body_run.bold = False
        body_run.font.size = Pt(10)

    risk_row = table.add_row()
    risk_cell = risk_row.cells[0].merge(risk_row.cells[5])
    risk_items = data.get("remarks_risk", [])
    risk_cell.paragraphs[0].clear()
    p = risk_cell.paragraphs[0]
    label_run = p.add_run("Remarks / Risk Points\n")
    label_run.bold = True
    label_run.font.size = Pt(10)
    for i, r in enumerate(risk_items):
        body_run = p.add_run(f"{i+1}. {r}\n")
        body_run.bold = False
        body_run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_report_docx(data: dict) -> bytes:
    """Generate the Bid Intelligence Report — clean 2-3 page document"""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    today = datetime.today().strftime("%d %B %Y")

    def heading(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
        if level == 1:
            run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x40)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def body(text):
        p = doc.add_paragraph(text)
        if p.runs:
            p.runs[0].font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)
        return p

    def bullet_list(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(str(item))
            run.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BID INTELLIGENCE REPORT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x40)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("MeraPath Education Limited")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = date_p.add_run(f"Generated: {today}")
    run3.font.size = Pt(10)

    doc.add_paragraph()

    # Section 1 — Tender Overview
    heading("1. Tender Overview")
    kv_table = doc.add_table(rows=7, cols=2)
    kv_table.style = 'Table Grid'
    rows_data = [
        ("Project Name",    data.get("project_name", "")),
        ("Department",      data.get("department_name", "")),
        ("Estimated Cost",  data.get("estimated_cost", "")),
        ("Bid Number",      data.get("bid_number", "")),
        ("Start Date",      data.get("start_date", "")),
        ("End Date",        data.get("end_date", "")),
        ("State / Region",  f"{data.get('state','')} — {data.get('region','')}"),
    ]
    for i, (k, v) in enumerate(rows_data):
        kv_table.cell(i, 0).text = k
        kv_table.cell(i, 1).text = str(v)
        if kv_table.cell(i, 0).paragraphs[0].runs:
            kv_table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Section 2 — Tender Jist
    heading("2. Summary")
    body(data.get("jist", ""))

    # Section 3 — Scope of Work
    heading("3. Scope of Work")
    bullet_list(data.get("scope_of_work", []))

    # Section 4 — Eligibility
    heading("4. Minimum Eligibility Criteria")
    bullet_list(data.get("eligibility_criteria", []))

    # Section 5 — Match Engine
    heading("5. MeraPath Match Engine")
    match = data.get("match_engine", {})
    criteria_labels = {
        "pmgdisha": "PMGDISHA Experience",
        "jjm":      "Jal Jeevan Mission",
        "lakhpati": "Lakhpati Didi Experience",
        "ndlm":     "NDLM Experience",
        "pmay":     "PMAY-G Experience",
        "iso":      "ISO Certifications",
        "msme":     "MSME Status",
        "geographic":"Geographic Reach",
        "scale":    "Beneficiary Scale",
    }
    match_table = doc.add_table(rows=len(criteria_labels)+1, cols=3)
    match_table.style = 'Table Grid'
    headers = ["Criteria", "Status", "MeraPath Credential"]
    details = {
        "pmgdisha":  "1,59,547 beneficiaries",
        "jjm":       "1,34,627 trainees",
        "lakhpati":  "65,186 beneficiaries",
        "ndlm":      "41,873 beneficiaries",
        "pmay":      "26,720 trainees",
        "iso":       "ISO 9001, 14001, 45001, 21001, 22000",
        "msme":      "MSME registered",
        "geographic":"14+ states, 3000+ panchayats",
        "scale":     "4.2 lakh+ beneficiaries",
    }
    for j, h in enumerate(headers):
        match_table.cell(0, j).text = h
        if match_table.cell(0, j).paragraphs[0].runs:
            match_table.cell(0, j).paragraphs[0].runs[0].bold = True
    for i, (key, label) in enumerate(criteria_labels.items(), 1):
        status = "✅ Match" if match.get(key) else "❌ Not applicable"
        match_table.cell(i, 0).text = label
        match_table.cell(i, 1).text = status
        match_table.cell(i, 2).text = details.get(key, "")

    doc.add_paragraph()

    # Section 6 — Score
    heading("6. Bid Strength Assessment")
    score = data.get("bid_strength_score", 0)
    prob  = data.get("win_probability", 0)

    score_table = doc.add_table(rows=2, cols=2)
    score_table.style = 'Table Grid'
    score_table.cell(0, 0).text = "Bid Strength Score"
    score_table.cell(0, 1).text = f"{score} / 100"
    score_table.cell(1, 0).text = "Win Probability"
    score_table.cell(1, 1).text = f"{prob}%"
    for row in score_table.rows:
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
        if row.cells[1].paragraphs[0].runs:
            row.cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    body(data.get("score_reasoning", ""))

    # Section 7 — Suggested Proposal Structure
    heading("7. Suggested Proposal Structure")
    bullet_list(data.get("proposal_sections", []))

    # Section 8 — Remarks
    heading("8. Remarks and Risk Points")
    bullet_list(data.get("remarks_risk", []))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if __name__ == "__main__":
    # Test data for checking DOCX format
    data = {
        "bid_number": "TEST/001",
        "ra_option": "No",
        "start_date": "1 June 2025",
        "end_date": "25 June 2025, 5:00 PM",
        "department_name": "Test Department, New Delhi — 110001",
        "estimated_cost": "₹2.5 Crore",
        "epbg": "₹5 Lakh",
        "processing_fee": "₹1,000",
        "tender_fee": "₹2,000",
        "appraisal_fee": "Not mentioned",
        "project_name": "Test Project for DOCX Format Check",
        "scope_of_work": ["Point 1 of scope", "Point 2 of scope", "Point 3 of scope"],
        "eligibility_criteria": ["Criterion 1", "Criterion 2", "Criterion 3"],
        "documents_required": ["Document 1", "Document 2"],
        "remarks_risk": ["Risk point 1", "Risk point 2"],
        "jist": "This is a test tender to verify the DOCX format is correct.",
        "state": "Delhi",
        "region": "Central Delhi",
        "category": "skill_development",
        "portal": "GeM",
        "match_engine": {
            "pmgdisha": True, "jjm": False, "lakhpati": True,
            "ndlm": True, "pmay": False, "iso": True,
            "msme": True, "geographic": True, "scale": True
        },
        "bid_strength_score": 82,
        "win_probability": 74,
        "score_reasoning": "Strong match on most criteria.",
        "proposal_sections": ["Section 1: Executive Summary", "Section 2: Methodology"]
    }

    top_bytes    = generate_top_docx(data)
    report_bytes = generate_report_docx(data)

    with open("test_TOP.docx", "wb") as f:
        f.write(top_bytes)
    with open("test_Report.docx", "wb") as f:
        f.write(report_bytes)

    print("✅ TOP.docx generated")
    print("✅ Report.docx generated")